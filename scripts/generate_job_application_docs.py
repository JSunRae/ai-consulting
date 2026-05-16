from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
RESEARCH_DIR = BASE_DIR / "Research and Documentation"
APPLICATIONS_DIR = RESEARCH_DIR / "Applications"
BASE_VARIANTS_DIR = APPLICATIONS_DIR / "Base Variants"

RESEARCH_DIR.mkdir(parents=True, exist_ok=True)
APPLICATIONS_DIR.mkdir(parents=True, exist_ok=True)
BASE_VARIANTS_DIR.mkdir(parents=True, exist_ok=True)

DATE_LINE = "12 May 2026"

NAME = "Jason Rae"
CONTACT = (
    "Stuttgart, Germany | Jason_C_Rae@Outlook.com | jasonrae.ai | "
    "linkedin.com/in/jason-c-rae | github.com/JSunRae"
)
PRIMARY = RGBColor(10, 25, 47)
TEXT = RGBColor(50, 50, 50)
SUBTLE = RGBColor(90, 90, 90)


@dataclass(frozen=True)
class VariantConfig:
    key: str
    file_stem: str
    subtitle: str
    target_role: str
    summary: str
    alignment_rows: list[tuple[str, str]]
    stack: str
    extra_current_bullets: list[str]
    extra_previous_bullets: list[str]
    cover_template_stem: str
    cover_template_title: str
    cover_template_paragraphs: list[str]


@dataclass(frozen=True)
class RoleConfig:
    key: str
    company: str
    role: str
    lane: str
    variant_key: str
    role_url: str
    checked_date: str
    location_model: str
    source: str
    fit: str
    comp_confidence: str
    work_auth_risk: str
    stretch_type: str
    action: str
    status: str
    packet_stem: str
    opening_reason: str
    alignment_rows: list[tuple[str, str]]
    strongest_alignments: list[str]
    risks: list[str]
    handling: list[str]
    proof_points: list[str]
    role_summary_tail: str
    role_specific_story: str
    form_answer_deltas: list[str]
    blocker_note: str
    short_answer: str


def set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = TEXT
    for style_name in ["Title", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.color.rgb = PRIMARY
    doc.styles["Title"].font.size = Pt(20)
    doc.styles["Heading 1"].font.size = Pt(12)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(10.5)
    doc.styles["Heading 2"].font.bold = True


def add_header(doc: Document, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(NAME)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.color.rgb = SUBTLE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CONTACT)
    run.font.size = Pt(9)
    run.font.color.rgb = SUBTLE


def add_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.font.color.rgb = PRIMARY


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    p.add_run(text)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(f"• {item}")
        run.font.size = Pt(10)


def add_role_block(
    doc: Document,
    title: str,
    company: str,
    location: str,
    dates: str,
    bullets: list[str],
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    left = p.add_run(f"{title} | {company}")
    left.bold = True
    left.font.color.rgb = PRIMARY
    p.add_run(f" | {location} | {dates}")
    add_bullets(doc, bullets)


def add_highlights_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    hdr[0].text = "Alignment Area"
    hdr[1].text = "Evidence"
    for cell in hdr:
        set_cell_shading(cell, "DCE6F1")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = PRIMARY
                run.font.size = Pt(9)
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
    for row in table.rows[1:]:
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(5.7)
        for idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.size = Pt(9.2)
                    if idx == 0:
                        run.bold = True
                        run.font.color.rgb = PRIMARY


def build_resume(
    path: Path,
    subtitle: str,
    target_role: str,
    summary: str,
    rows: list[tuple[str, str]],
    experience: list[dict[str, object]],
    earlier_highlights: list[str],
    education: str,
    credentials: str,
    stack: str,
) -> None:
    doc = Document()
    set_page(doc)
    style_doc(doc)
    add_header(doc, subtitle)
    add_section(doc, "Target Role")
    add_body(doc, target_role)
    add_section(doc, "Executive Profile")
    add_body(doc, summary)
    add_section(doc, "Role Alignment")
    add_highlights_table(doc, rows)
    add_section(doc, "Professional Experience")
    for job in experience:
        add_role_block(
            doc,
            str(job["title"]),
            str(job["company"]),
            str(job["location"]),
            str(job["dates"]),
            list(job["bullets"]),
        )
    add_section(doc, "Earlier Career Highlights")
    add_bullets(doc, earlier_highlights)
    add_section(doc, "Education")
    add_body(doc, education)
    add_section(doc, "Selected Credentials")
    add_body(doc, credentials)
    add_section(doc, "Technical and Operating Stack")
    add_body(doc, stack)
    doc.save(path)


def build_cover_letter(
    path: Path,
    subtitle: str,
    date_line: str,
    company: str,
    greeting: str,
    paragraphs: list[str],
) -> None:
    doc = Document()
    set_page(doc)
    style_doc(doc)
    add_header(doc, subtitle)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(date_line)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(company)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(greeting)

    for para in paragraphs:
        add_body(doc, para)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run("Kind regards,")

    p = doc.add_paragraph()
    p.add_run(NAME)

    doc.save(path)


def build_recruiter_one_pager(
    path: Path,
    subtitle: str,
    summary: str,
    target_roles: list[str],
    proof_points: list[str],
    search_mandate_fit: list[str],
    stack: str,
) -> None:
    doc = Document()
    set_page(doc)
    style_doc(doc)
    add_header(doc, subtitle)
    add_section(doc, "Executive Profile")
    add_body(doc, summary)
    add_section(doc, "Target Role Families")
    add_bullets(doc, target_roles)
    add_section(doc, "Signature Proof Points")
    add_bullets(doc, proof_points)
    add_section(doc, "Search-Mandate Fit")
    add_bullets(doc, search_mandate_fit)
    add_section(doc, "Technical and Operating Stack")
    add_body(doc, stack)
    doc.save(path)


def clone_experience(variant: VariantConfig) -> list[dict[str, object]]:
    experience = deepcopy(COMMON_EXPERIENCE)
    if variant.extra_current_bullets:
        experience[0]["bullets"].extend(variant.extra_current_bullets)
    if variant.extra_previous_bullets:
        experience[1]["bullets"].extend(variant.extra_previous_bullets)
    return experience


def role_resume_summary(variant: VariantConfig, role: RoleConfig) -> str:
    return f"{variant.summary} {role.role_summary_tail}".strip()


def compose_role_cover_letter(
    variant: VariantConfig,
    role: RoleConfig,
) -> list[str]:
    first_outcome = role.proof_points[0]
    second_outcome = role.proof_points[1]
    return [
        (
            f"I am applying for the {role.role} role at {role.company} because "
            f"{role.opening_reason}"
        ),
        (
            "I currently own EMEA commercial analytics at Medline across pricing, "
            "forecasting, margin, CRM analytics, incentives, and executive business "
            "reviews. My background is unusual in a useful way: I started in "
            "commercial leadership, then moved into senior analytics and applied AI "
            "work. That means I understand both the operating pressure behind growth, "
            "forecasting, and decision-quality calls and the workflow discipline "
            "needed to support them reliably."
        ),
        role.role_specific_story,
        (
            f"The measurable outcomes matter. I delivered {first_outcome}, "
            f"{second_outcome}, and have repeatedly operated across Sales, Finance, "
            "BI, IT, and leadership to improve the quality and speed of decisions. "
            "That cross-functional operating model is one of the clearest reasons I "
            f"am relevant to this {role.lane.lower()} opening."
        ),
        (
            f"I would welcome the opportunity to discuss how I could help "
            f"{role.company} strengthen this capability with a practical, hands-on "
            "leadership style centered on measurable value creation."
        ),
    ]


def render_packet_markdown(
    role: RoleConfig,
    variant: VariantConfig,
    resume_filename: str,
    cover_filename: str,
) -> str:
    rows = "\n".join(
        f"| {left} | {right} |"
        for left, right in [
            ("Checked", role.checked_date),
            ("Lane", role.lane),
            ("Location / model", role.location_model),
            ("Status", role.status),
            ("Action", role.action),
            ("Fit", role.fit),
            ("Comp confidence", role.comp_confidence),
            ("Stretch type", role.stretch_type),
            ("Work auth risk", role.work_auth_risk),
        ]
    )
    strongest = "\n".join(f"- {item}" for item in role.strongest_alignments)
    risks = "\n".join(f"- {item}" for item in role.risks)
    handling = "\n".join(f"- {item}" for item in role.handling)
    proof_points = "\n".join(f"- `{item}`" for item in role.proof_points)
    form_deltas = "\n".join(f"- {item}" for item in role.form_answer_deltas)

    return f"""# {role.company} - {role.role}

Prepared: `{DATE_LINE}`  
Role URL: `{role.role_url}`

## Packet status

| Field | Value |
| --- | --- |
{rows}

## Base variant used

- Resume base: `Applications/Base Variants/{variant.file_stem}.docx`
- Cover-letter base: `Applications/Base Variants/{variant.cover_template_stem}.docx`
- Role-specific resume: `Applications/{resume_filename}`
- Role-specific cover letter: `Applications/{cover_filename}`

## Fit assessment

Overall fit: `{role.fit}`

Strongest alignments:

{strongest}

Main risks:

{risks}

How to handle the risks:

{handling}

## Tailored executive summary

{role_resume_summary(variant, role)}

## Priority proof points to keep

{proof_points}

## Form-answer deltas

{form_deltas}

## Blocker note

{role.blocker_note}

## Short answer draft

```text
{role.short_answer}
```
"""


COMMON_EXPERIENCE: list[dict[str, object]] = [
    {
        "title": "Senior Data Analyst EMEA",
        "company": "Medline",
        "location": "Stuttgart, Germany",
        "dates": "Jan 2026 - Present",
        "bullets": [
            "Own EMEA commercial analytics across gross margin walks, pricing intelligence, forecasting, CRM analytics, opportunity identification, incentives, and executive business reviews.",
            "Build decision-ready workflows across Power BI, Python, SQL, machine learning, and controlled LLM systems where they improve auditability and speed.",
            "Partner across Sales, Finance, BI, IT, and leadership to improve margin visibility, forecast quality, pricing interpretation, CRM trust, and decision speed.",
            "Translate complex commercial data into executive narratives connecting revenue, margin, price, volume, mix, FX, and customer behavior.",
        ],
    },
    {
        "title": "Senior Data Analyst Europe",
        "company": "Medline",
        "location": "Germany, Hybrid",
        "dates": "Oct 2022 - Jan 2026",
        "bullets": [
            "Delivered EUR 1.3M UK sales growth in 2023, approximately 4x above target, by combining commercial execution, analytics, and AI-enabled opportunity identification while reporting to the SVP Sales Europe.",
            "Corrected 10+ years of FX methodology errors and rebuilt Price / Volume / FX decomposition logic with correct interaction effects.",
            "Designed an enterprise-grade bonus payment system and led cross-functional problem solving across Sales, Finance, BI, and IT to improve trust in commercial performance data.",
            "Led customer-level integration across two acquisitions, coordinating regional analysts and country teams to reconcile credits, refactor customer/account data, preserve continuity, and restore reliable territory measurement.",
        ],
    },
    {
        "title": "Inside Sales Department Manager / BDM UK",
        "company": "Medline",
        "location": "Warrington, UK",
        "dates": "Oct 2020 - Sep 2022",
        "bullets": [
            "Launched a new UK Inside Sales department and delivered EUR 3M sales growth against an initial target of approximately EUR 300k through data-driven segmentation, targeting, and execution.",
            "Built analytics for customer segmentation, forecasting, market analysis, and backorder management to prioritize accounts and guide commercial strategy.",
            "Identified 120 new distributors and secured 33 new accounts under pandemic and GDPR constraints.",
        ],
    },
]

EARLIER_HIGHLIGHTS = [
    "Inside Sales Department Manager Australia: drove a 4% to 57% growth turnaround across 38 customer markets.",
    "Account Manager Australia NSW and ACT: grew territory revenue from USD 1.4M to USD 4.2M and achieved 35% profitability.",
    "Pilot, Royal Australian Air Force: developed disciplined systems thinking, rapid learning, and high-pressure decision-making.",
]

EDUCATION = "Bachelor of Technology (Aviation), Australian Defence Force Academy (UNSW), 2008."
SELECTED_CREDENTIALS = (
    "Python | Machine Learning | Microsoft Excel | VBA | PSS Selling | "
    "The Challenger Sale"
)

VARIANTS: dict[str, VariantConfig] = {
    "analytics_leadership": VariantConfig(
        key="analytics_leadership",
        file_stem="Jason_Rae_Analytics_Leadership_Resume",
        subtitle="Commercial Analytics & Applied AI Leader | Analytics Leadership Target",
        target_role=(
            "Head / Director of Analytics | Head / Director of Data & Analytics | "
            "Head of Commercial Analytics | KPI Governance | Executive Decision Systems"
        ),
        summary=(
            "Commercial analytics and applied AI leader with 13+ years across "
            "commercial leadership and analytics, including pricing, forecasting, "
            "margin analysis, CRM governance, acquisition integration, and data "
            "transformation. Proven record of building trusted decision systems, "
            "correcting broken measurement logic, and translating fragmented data "
            "into executive-ready insights across Sales, Finance, BI, IT, and "
            "leadership teams. Combines strategic analytics ownership with hands-on "
            "delivery in Python, SQL, Power BI, machine learning, and governed LLM "
            "workflows. Most relevant where a business needs stronger KPI governance, "
            "analytics maturity, cross-functional operating rhythms, and practical "
            "AI-enabled decision support."
        ),
        alignment_rows=[
            ("KPI governance", "Corrected 10+ years of FX and reporting methodology issues and rebuilt PVFX logic to restore trust in commercial decisions."),
            ("Executive decision systems", "Own EMEA commercial analytics across pricing, forecasting, CRM analytics, incentives, and executive business reviews."),
            ("Commercial impact", "Delivered EUR 3M sales growth against an initial target of approximately EUR 300k and EUR 1.3M UK growth in 2023 at roughly 4x target."),
            ("Cross-functional leadership", "Operate across Sales, Finance, BI, IT, and country teams, including customer-level integration across two acquisitions."),
            ("Applied AI", "Build controlled workflows across Python, SQL, Power BI, ML, and LLM systems where they improve auditability and speed."),
        ],
        stack=(
            "Python | SQL | Power BI | Machine Learning | LLM Workflows | KPI Governance | "
            "Executive Reporting | CRM Analytics | Forecasting | Pricing | Margin Analysis | "
            "Data Quality | Cross-Functional Leadership"
        ),
        extra_current_bullets=[
            "Lead applied AI and analytics enablement by translating technical capability into practical commercial use cases, training approaches, and governance models."
        ],
        extra_previous_bullets=[
            "Led European contribution to a global applied AI initiative, building a training model to help teams adopt AI safely across real workflows."
        ],
        cover_template_stem="Jason_Rae_Analytics_Leadership_Cover_Letter_Template",
        cover_template_title="Cover Letter Template | Analytics Leadership",
        cover_template_paragraphs=[
            "I am targeting Head and Director-level analytics leadership roles where the mandate goes beyond reporting and into KPI governance, executive decision systems, and measurable operating change.",
            "My background combines commercial leadership with hands-on analytics and applied AI delivery. I currently own EMEA commercial analytics at Medline across pricing, forecasting, margin, CRM analytics, incentives, and executive business reviews.",
            "Across recent roles, I have corrected broken reporting logic, rebuilt trust in performance data, and created decision-ready workflows across Python, SQL, Power BI, machine learning, and governed LLM systems.",
            "The measurable outcomes matter. I delivered EUR 3M sales growth against an initial target of approximately EUR 300k, generated EUR 1.3M UK growth in 2023 at roughly 4x target, and repeatedly improved executive decision quality by strengthening how data is interpreted and used.",
            "I would bring a practical, hands-on leadership style centered on KPI clarity, data trust, cross-functional influence, and measurable value creation.",
        ],
    ),
    "product_saas_analytics": VariantConfig(
        key="product_saas_analytics",
        file_stem="Jason_Rae_Product_SaaS_Analytics_Resume",
        subtitle="Commercial Analytics & Applied AI Leader | Product / SaaS Analytics Target",
        target_role=(
            "Director / Head of Product Analytics | Product & Business Analytics | "
            "SaaS Metrics | Funnel Analytics | Self-Serve Analytics"
        ),
        summary=(
            "Commercial analytics and applied AI leader with 13+ years across "
            "commercial leadership, analytics, pricing, forecasting, CRM governance, "
            "and executive decision support. Strongest where a SaaS or digital product "
            "business needs clearer funnel visibility, more trusted KPI systems, "
            "better self-serve analytics, and a player-coach leader who can connect "
            "customer behavior to business performance. Brings hands-on delivery in "
            "SQL, Python, Power BI, machine learning, and governed LLM workflows, "
            "plus a strong cross-functional operating style across commercial, finance, "
            "product-adjacent, and executive stakeholders."
        ),
        alignment_rows=[
            ("Measurement strategy", "Own EMEA commercial analytics across pricing, forecasting, CRM analytics, incentives, and executive business reviews."),
            ("Funnel and conversion logic", "Translate complex commercial data into narratives connecting opportunity quality, conversion, revenue, and customer behavior."),
            ("Self-serve analytics", "Build decision-ready workflows across Power BI, Python, SQL, machine learning, and controlled LLM systems."),
            ("Growth impact", "Delivered EUR 3M sales growth against an initial target of approximately EUR 300k and EUR 1.3M UK growth in 2023 at roughly 4x target."),
            ("Cross-functional translation", "Operate across Sales, Finance, BI, IT, and leadership to improve decision speed and reporting trust."),
        ],
        stack=(
            "SQL | Python | Power BI | Machine Learning | LLM Workflows | Funnel Analytics | "
            "Executive Reporting | CRM Analytics | Forecasting | Pricing | Margin Analysis | "
            "KPI Governance | Cross-Functional Leadership"
        ),
        extra_current_bullets=[
            "Translate complex operating data into executive narratives linking pipeline quality, conversion, revenue, margin, and customer behavior."
        ],
        extra_previous_bullets=[
            "Built analytics-driven opportunity identification workflows that linked customer behavior to revenue actions and executive prioritization."
        ],
        cover_template_stem="Jason_Rae_Product_SaaS_Analytics_Cover_Letter_Template",
        cover_template_title="Cover Letter Template | Product / SaaS Analytics",
        cover_template_paragraphs=[
            "I am targeting product and SaaS analytics leadership roles where analytics needs to shape product priorities, funnel performance, retention, and executive decisions rather than operate as a reporting service.",
            "My background combines direct commercial operating experience with hands-on analytics and applied AI delivery. I currently own EMEA commercial analytics at Medline across pricing, forecasting, margin, CRM analytics, incentives, and executive business reviews.",
            "Across recent roles, I have built trusted KPI systems, improved reporting integrity, and translated complex operating data into clear actions for leadership, commercial teams, and technical stakeholders.",
            "The outcomes are measurable. I delivered EUR 3M sales growth against an initial target of approximately EUR 300k, generated EUR 1.3M UK growth in 2023 at roughly 4x target, and repeatedly improved decision quality by making commercial signals clearer and more actionable.",
            "I would bring a practical player-coach style centered on measurement clarity, self-serve analytics, strong stakeholder translation, and commercially useful insight.",
        ],
    ),
    "ai_transformation_consulting": VariantConfig(
        key="ai_transformation_consulting",
        file_stem="Jason_Rae_AI_Transformation_Consulting_Resume",
        subtitle="Commercial Analytics & Applied AI Leader | AI Transformation Target",
        target_role=(
            "Head of AI | AI Transformation Lead | AI Consulting Director | "
            "Applied AI Productization | Governance | Executive Advisory"
        ),
        summary=(
            "Commercial analytics and applied AI leader with 13+ years across "
            "commercial leadership, analytics, pricing, forecasting, margin analysis, "
            "CRM governance, acquisition integration, and data transformation. Brings "
            "a rare combination of commercial operating judgment, hands-on machine "
            "learning depth, and practical AI workflow delivery. Strongest where AI "
            "needs to move from concept to governed business use cases, with clear "
            "stakeholder alignment, executive visibility, and measurable value "
            "creation. Experienced in Python, SQL, Power BI, machine learning, "
            "TensorFlow since 2016, and controlled LLM workflows, with a strong "
            "cross-functional operating style across business and technical teams."
        ),
        alignment_rows=[
            ("AI enablement", "Led European contribution to a global applied AI initiative and built training paths for safer AI adoption across real workflows."),
            ("Business translation", "Translate technical capability into practical use cases, executive-ready communication, and governed rollout paths."),
            ("Cross-functional delivery", "Operate across Sales, Finance, BI, IT, and leadership in complex multi-country environments."),
            ("Hands-on depth", "Build in Python, SQL, TensorFlow, machine learning, and governed LLM workflows rather than treating AI as abstract strategy."),
            ("Commercial relevance", "Anchor AI work in forecasting, pricing, margin, CRM, and executive decision support rather than novelty demos."),
        ],
        stack=(
            "Python | SQL | TensorFlow | Machine Learning | LLM Workflows | AI Enablement | "
            "Workflow Automation | Executive Reporting | Forecasting | Pricing | CRM Analytics | "
            "Cross-Functional Leadership | Governance"
        ),
        extra_current_bullets=[
            "Lead applied AI and analytics enablement by translating technical capability into practical commercial use cases, training approaches, and governance models."
        ],
        extra_previous_bullets=[
            "Led European contribution to a global applied AI initiative, building a training model to help teams adopt AI safely across real workflows."
        ],
        cover_template_stem="Jason_Rae_AI_Transformation_Cover_Letter_Template",
        cover_template_title="Cover Letter Template | AI Transformation / Consulting",
        cover_template_paragraphs=[
            "I am targeting AI leadership and AI consulting roles where the mandate is practical adoption, productization, governance, and measurable business value rather than pure research or infrastructure ownership.",
            "My background combines commercial leadership, senior analytics work, and hands-on machine learning and workflow delivery. I currently own EMEA commercial analytics at Medline across pricing, forecasting, margin, CRM analytics, incentives, and executive business reviews.",
            "One of the strongest parts of my recent experience has been translating technical AI capability into workflows people can actually adopt, with training, governance, and business alignment built in from the start.",
            "The measurable outcomes matter. I delivered EUR 3M sales growth against an initial target of approximately EUR 300k, generated EUR 1.3M UK growth in 2023 at roughly 4x target, and have repeatedly linked analytics and AI work to real operating decisions.",
            "I would bring a practical leadership style, strong stakeholder management, and enough technical depth to be useful both strategically and hands-on.",
        ],
    ),
}

ROLES: list[RoleConfig] = [
    RoleConfig(
        key="taxdome",
        company="TaxDome",
        role="Head of Analytics",
        lane="Analytics leadership",
        variant_key="analytics_leadership",
        role_url="https://careers.taxdome.com/v/164918-head-of-analytics",
        checked_date="2026-05-08",
        location_model="Fully remote global",
        source="Direct company page",
        fit="High",
        comp_confidence="Moderate-high; likely above EUR 125k total if level is genuine Head scope",
        work_auth_risk="Low",
        stretch_type="Manageable stretch on product analytics and data engineering language",
        action="Monitor submitted application",
        status="Submitted",
        packet_stem="TaxDome_Head_of_Analytics",
        opening_reason="the mandate is much broader than reporting. It is about building the measurement systems, governance discipline, and cross-functional operating model that help a company make smarter, faster decisions. That is where my experience is strongest.",
        alignment_rows=[
            ("Analytics strategy", "Own EMEA commercial analytics across pricing, forecasting, CRM analytics, incentives, and executive business reviews."),
            ("Governance", "Corrected 10+ years of FX and reporting methodology issues and rebuilt PVFX logic to restore trust in commercial decisions."),
            ("Executive reporting", "Translate complex commercial data into executive narratives linking revenue, margin, price, volume, mix, FX, and customer behavior."),
            ("AI-enabled analytics", "Build controlled workflows across Python, SQL, Power BI, ML, and LLM systems where they improve auditability and speed."),
            ("Cross-functional leadership", "Operate across Sales, Finance, BI, IT, and country teams, including customer-level integration across two acquisitions."),
        ],
        strongest_alignments=[
            "Analytics leadership tied to business outcomes, KPI definition, and executive decision quality.",
            "Cross-functional work across Sales, Finance, BI, IT, and leadership.",
            "Governance, reporting trust, and data-quality correction work.",
            "Hands-on Python, SQL, and BI depth alongside strategic leadership narrative.",
            "Practical applied AI story rather than novelty AI positioning.",
        ],
        risks=[
            "Role expects stronger direct product analytics signal than the public materials initially showed.",
            "Role spans data engineering plus product analytics tooling such as BigQuery, Looker, Amplitude, GA4, HubSpot, and Fivetran.",
            "Strong SaaS metric vocabulary is required: ARR, activation, retention, NDR, CSAT.",
        ],
        handling=[
            "Lead with KPI ownership, governance, and cross-functional measurement systems.",
            "Translate commercial analytics work into product and business decision-system language.",
            "Be explicit in interviews that the value is building trusted measurement frameworks and operating cadences, not only running analyses.",
        ],
        proof_points=[
            "Own EMEA commercial analytics across pricing, forecasting, CRM analytics, incentives, and executive business reviews",
            "corrected 10+ years of FX and reporting methodology issues affecting trust in commercial decisions",
            "delivered EUR 3M sales growth against an initial target of approximately EUR 300k",
            "delivered EUR 1.3M UK growth in 2023 at roughly 4x target",
        ],
        role_summary_tail="Most relevant where a business needs stronger KPI governance, analytics maturity, cross-functional operating rhythms, and practical AI-enabled decision support.",
        role_specific_story="Recent work has focused on turning fragmented data into trusted decision systems. I have corrected long-running FX and reporting methodology issues, rebuilt Price / Volume / FX logic so leadership could trust the story behind performance, and built decision-ready workflows across Python, SQL, Power BI, machine learning, and governed LLM systems. I also led customer-level integration work across two acquisitions, coordinating analysts and country teams to preserve continuity and restore reliable performance measurement.",
        form_answer_deltas=[
            "Use the analytics leadership resume and role-specific cover letter already prepared.",
            "Lean into analytics maturity, KPI governance, and AI-assisted analytics rather than trying to mimic a pure product-analytics profile.",
        ],
        blocker_note="No blocker. Application already submitted and should only be updated if TaxDome asks for refreshed materials.",
        short_answer="TaxDome is attractive because the role is not limited to dashboard ownership. It combines analytics strategy, KPI governance, data foundations, AI-assisted analytics, and cross-functional influence across Product, Engineering, Revenue, and Operations.",
    ),
    RoleConfig(
        key="joor",
        company="JOOR",
        role="Head of Analytics",
        lane="Analytics leadership",
        variant_key="analytics_leadership",
        role_url="https://jobs.ashbyhq.com/joor/3410ed7b-6161-444f-81ea-fd48f5a53860",
        checked_date="2026-05-08",
        location_model="Spain remote listing / Europe context",
        source="Direct company page",
        fit="Medium-high",
        comp_confidence="Moderate; title and scope likely above EUR 125k total if Europe-remote band is competitive",
        work_auth_risk="Low-medium",
        stretch_type="Manageable stretch on marketplace and product analytics vocabulary",
        action="Monitor submitted application",
        status="Submitted",
        packet_stem="JOOR_Head_of_Analytics",
        opening_reason="it combines the parts of analytics leadership where I add the most value: business intelligence, KPI governance, scalable reporting, and turning complex data into clear actions for leadership and cross-functional teams.",
        alignment_rows=[
            ("Business intelligence", "Build executive-ready reporting and analytical workflows across Power BI, Python, SQL, ML, and controlled LLM systems."),
            ("KPI visibility", "Own EMEA commercial analytics across pricing, forecasting, CRM analytics, incentives, and executive business reviews."),
            ("Data integrity", "Corrected long-running FX and reporting logic issues that were distorting commercial decisions."),
            ("Revenue impact", "Delivered EUR 3M sales growth against an initial target of approximately EUR 300k and EUR 1.3M UK growth in 2023 at roughly 4x target."),
            ("Cross-functional work", "Partner across Sales, Finance, BI, IT, and country teams to improve trust in operating data and decision speed."),
        ],
        strongest_alignments=[
            "Revenue-linked analytics leadership with strong executive reporting depth.",
            "KPI governance and data trust story that maps well to marketplace operating complexity.",
            "Hands-on SQL, Python, Power BI, and workflow automation credibility.",
            "Cross-functional operating style across commercial, finance, BI, and leadership stakeholders.",
        ],
        risks=[
            "Sharper product and marketplace framing is needed than the base commercial analytics story provides by default.",
            "Some direct marketplace, B2B SaaS, and e-commerce vocabulary has to be bridged carefully.",
        ],
        handling=[
            "Lead with business intelligence, scalable reporting, and measurable decision support.",
            "Translate commercial analytics work into revenue, customer, and network-effect language where truthful.",
            "Avoid overstating direct marketplace depth and instead emphasize transferable measurement leadership.",
        ],
        proof_points=[
            "Own EMEA commercial analytics across pricing, forecasting, CRM analytics, incentives, and executive business reviews",
            "delivered EUR 1.3M UK growth in 2023 at roughly 4x target",
            "delivered EUR 3M sales growth against an initial target of approximately EUR 300k",
            "corrected long-running FX and reporting logic issues affecting trust in commercial decisions",
        ],
        role_summary_tail="Strong fit for analytics leadership roles that require business intelligence, data quality, scalable reporting, and revenue-linked analytics strategy.",
        role_specific_story="Across recent roles, I have built executive-ready reporting and analytical workflows, corrected long-running data and reporting issues that were distorting commercial decisions, and partnered across Sales, Finance, BI, and IT to improve trust in performance data. I have also led customer-level integration across two acquisitions, coordinating analysts and country teams to preserve continuity and maintain reliable measurement.",
        form_answer_deltas=[
            "Use the analytics leadership variant with revenue and business-intelligence emphasis.",
            "Keep marketplace and product language selective and evidence-based.",
        ],
        blocker_note="No blocker. Application already submitted manually using the prepared packet.",
        short_answer="JOOR is attractive because the role sits at the intersection of business intelligence, KPI governance, and scalable revenue analytics rather than reporting in isolation.",
    ),
    RoleConfig(
        key="smallpdf",
        company="smallpdf",
        role="Head of Data Analytics",
        lane="Analytics leadership",
        variant_key="analytics_leadership",
        role_url="https://jobs.ashbyhq.com/smallpdf/220df7f4-bd13-42a6-8b02-f9771961fa14",
        checked_date="2026-05-08",
        location_model="Remote",
        source="Direct company page / Ashby",
        fit="Medium-high",
        comp_confidence="Moderate-high; role scope likely worth the floor if work authorization fit existed",
        work_auth_risk="High",
        stretch_type="Good content fit but blocked on form wording",
        action="Keep blocked unless exact work-authorization wording becomes truthful",
        status="Blocked",
        packet_stem="smallpdf_Head_of_Data_Analytics",
        opening_reason="the role is framed in exactly the right way: not as dashboard production, but as a business leadership mandate where analytics should shape outcomes, strengthen decisions, and build a more insight-driven culture.",
        alignment_rows=[
            ("Business-first analytics", "Operate where analytics must shape outcomes rather than only produce dashboards or reports."),
            ("Insight generation", "Own EMEA commercial analytics across pricing, forecasting, CRM analytics, incentives, and executive business reviews."),
            ("Hands-on depth", "Build workflows across Python, SQL, Power BI, ML, and controlled LLM systems while staying close to the data."),
            ("Data trust", "Corrected long-running FX and reporting issues that were distorting commercial decisions."),
            ("Organizational change", "Help teams adopt analytics and AI more safely and practically through enablement and cross-functional coordination."),
        ],
        strongest_alignments=[
            "Business-first analytics leadership rather than dashboard production.",
            "Commercial analytics, KPI governance, and executive decision support.",
            "Strong SQL, Python, and BI orientation with hands-on analytical credibility.",
            "Experience changing how teams use data, not just supplying reports.",
            "Practical AI and analytics enablement story that fits the role's AI language.",
        ],
        risks=[
            "Stronger SaaS, subscription, and consumer behavior signal would help.",
            "Role mentions team scaling, SQL, dbt, and Looker specifically.",
            "The strongest public story is still commercial analytics rather than classic product analytics.",
        ],
        handling=[
            "Emphasize business-outcome orientation, proactive insight generation, and analytics culture change.",
            "Lead with player-coach style and the ability to test hypotheses directly in SQL and Python.",
            "Translate commercial analytics work into growth, retention, monetization, and behavioral-insight language where truthful.",
        ],
        proof_points=[
            "Own EMEA commercial analytics across pricing, forecasting, CRM analytics, incentives, and executive business reviews",
            "built decision-ready workflows across Power BI, Python, SQL, ML, and controlled LLM systems",
            "delivered EUR 1.3M UK growth in 2023 at roughly 4x target",
            "delivered EUR 3M sales growth against an initial target of approximately EUR 300k",
        ],
        role_summary_tail="Strong fit for analytics leadership roles that require business-first insight generation, data trust, organizational influence, and practical AI-enabled analytics.",
        role_specific_story="Across recent roles, I have turned fragmented data into trusted decision systems, corrected long-running reporting and FX methodology issues, rebuilt commercial logic leadership relied on, and created workflows across Python, SQL, Power BI, machine learning, and governed LLM systems that improved both speed and confidence. I have also helped teams adopt analytics and AI more safely and practically through training, cross-functional coordination, and clearer operating structures.",
        form_answer_deltas=[
            "No submission unless the work-authorization category can be answered exactly and truthfully.",
            "If the blocker is resolved, use the analytics leadership variant with business-first and culture-change framing.",
        ],
        blocker_note="Browser automation confirmed the form only accepts narrow passport or permit categories for Switzerland, Spain, Serbia, other EU/EEA countries, or no valid authorization. Leave blocked unless one of those exact categories is true.",
        short_answer="smallpdf is attractive because the role is explicitly about making analytics a business partner rather than a reporting service. That matches how I work best: building trusted KPI systems, improving decision quality, and helping leadership act on insight rather than just consume data.",
    ),
    RoleConfig(
        key="team_blue",
        company="team.blue",
        role="Head of AI (SaaS)",
        lane="AI transformation / consulting",
        variant_key="ai_transformation_consulting",
        role_url="https://careers.team.blue/jobs/6188773-head-of-ai-saas",
        checked_date="2026-05-08",
        location_model="Fully remote, multiple EU locations",
        source="Direct company page",
        fit="Medium",
        comp_confidence="Moderate-high; likely above floor if true Head-level scope",
        work_auth_risk="Low",
        stretch_type="Selective stretch on direct shipped AI product depth",
        action="Monitor submitted application",
        status="Submitted",
        packet_stem="team_blue_Head_of_AI_SaaS",
        opening_reason="it sits in a space where I can add differentiated value: translating AI capability into practical, governed business use cases across multiple brands, stakeholders, and operating realities.",
        alignment_rows=[
            ("AI enablement", "Led European contribution to a global applied AI initiative and built training paths for safer AI adoption across real workflows."),
            ("Business translation", "Translate technical capability into practical business use cases and executive-ready communication."),
            ("Cross-functional influence", "Operate across Sales, Finance, BI, IT, and leadership in complex multi-country environments."),
            ("Federated operations", "Led customer-level integration across two acquisitions, useful for multi-brand and post-merger coordination language."),
            ("Hands-on depth", "Build in Python, SQL, TensorFlow, machine learning, and governed LLM workflows rather than treating AI as abstract strategy."),
        ],
        strongest_alignments=[
            "AI enablement and workflow governance rather than AI theater.",
            "Cross-functional leadership across business and technical stakeholders.",
            "Useful experience for federated, multi-brand coordination language.",
            "Hands-on build credibility in TensorFlow, machine learning, and LLM workflows.",
        ],
        risks=[
            "Role wants stronger direct evidence of shipping AI products to end users.",
            "More explicit product roadmap and matrix influence examples help.",
        ],
        handling=[
            "Lead with AI enablement, safer rollout, and multi-stakeholder translation.",
            "Frame acquisition integration and regional coordination as relevant evidence for portfolio and brand alignment work.",
            "Stay grounded in business-value and governance language instead of claiming research depth.",
        ],
        proof_points=[
            "led European contribution to a global applied AI initiative with safer adoption and training paths",
            "own EMEA commercial analytics across pricing, forecasting, margin, CRM, incentives, and executive business reviews",
            "delivered EUR 1.3M UK growth in 2023 at roughly 4x target",
            "delivered EUR 3M sales growth against an initial target of approximately EUR 300k",
        ],
        role_summary_tail="Strongest where AI needs to move from concept to governed business use cases, with clear stakeholder alignment, executive visibility, and measurable value creation.",
        role_specific_story="One of the strongest parts of my recent experience was leading European contribution to a global applied AI initiative, where the challenge was not only technical possibility but safe adoption, enablement, and business integration. That required translating AI capability into workflows people could actually use, creating training and governance paths, and working across commercial and technical stakeholders with different priorities and levels of confidence.",
        form_answer_deltas=[
            "Use the AI transformation resume variant with AI enablement and product-roadmap language.",
            "Lead with business integration and governance rather than pure research or platform depth.",
        ],
        blocker_note="No blocker. Application already submitted and should only be revisited if the team asks for updated materials.",
        short_answer="team.blue is attractive because the role is about enabling AI product innovation across a federated SaaS portfolio, which fits my mix of practical AI adoption, cross-functional translation, and governance-minded delivery.",
    ),
    RoleConfig(
        key="checkmk",
        company="Checkmk",
        role="Head of Data & Analytics (m/f/d)",
        lane="Analytics leadership",
        variant_key="analytics_leadership",
        role_url="https://checkmk-gmbh.jobs.personio.de/job/2486975?display=en&hss_channel=lcp-11824415&language=en",
        checked_date="2026-05-12",
        location_model="Germany remote / Munich optional",
        source="Direct company page",
        fit="High",
        comp_confidence="High enough to pursue; prior internal estimate EUR 135k-170k total",
        work_auth_risk="Low",
        stretch_type="Manageable stretch on SaaS metrics and Tableau specificity",
        action="Apply next",
        status="Prepared",
        packet_stem="Checkmk_Head_of_Data_Analytics",
        opening_reason="the mandate goes well beyond reporting. It combines BI roadmap ownership, sales-funnel analytics, data governance, stakeholder translation, and hands-on strategic execution in a way that matches where I add the most value.",
        alignment_rows=[
            ("BI roadmap", "Own EMEA commercial analytics across pricing, forecasting, CRM analytics, incentives, and executive business reviews."),
            ("Stakeholder translation", "Translate complex commercial data into executive narratives linking revenue, margin, price, volume, mix, FX, and customer behavior."),
            ("Reporting trust", "Corrected 10+ years of FX and reporting methodology issues and rebuilt PVFX logic to restore decision confidence."),
            ("Hands-on leadership", "Build workflows across Python, SQL, Power BI, ML, and controlled LLM systems while staying close to the analytical work."),
            ("Sales-funnel relevance", "Built analytics for customer segmentation, forecasting, targeting, and account prioritization to guide commercial strategy."),
        ],
        strongest_alignments=[
            "Germany-remote compatibility with English as the core language.",
            "Clear BI roadmap, sales-funnel analytics, demand management, and stakeholder-translation remit.",
            "Strong fit to KPI governance, reporting trust, and structured intake of business questions.",
            "Leadership without requiring a pure research-ML background.",
        ],
        risks=[
            "The posting asks for deeper SaaS metric fluency across ARR, churn, CAC, and LTV than the base commercial analytics story naturally signals.",
            "Checkmk explicitly mentions Tableau Cloud, Postgres, and oversight of Kedro-based ETL pipelines.",
        ],
        handling=[
            "Lead with KPI governance, funnel visibility, and executive decision-system ownership.",
            "Position SQL fluency, analytics product ownership, and data-quality leadership clearly, without overstating Tableau specialization.",
            "Translate commercial pipeline, pricing, and CRM work into SaaS-funnel and revenue-quality language where truthful.",
        ],
        proof_points=[
            "own EMEA commercial analytics across pricing, forecasting, CRM analytics, incentives, and executive business reviews",
            "corrected 10+ years of FX and reporting methodology issues and rebuilt PVFX logic",
            "delivered EUR 3M sales growth against an initial target of approximately EUR 300k",
            "led customer-level integration across two acquisitions while preserving measurement quality",
        ],
        role_summary_tail="Most relevant where a SaaS company needs BI roadmap ownership, funnel visibility, metric trust, and a hands-on strategic leader who can improve both decision quality and adoption.",
        role_specific_story="Recent work has focused on building trusted decision systems, improving reporting integrity, and translating vague business questions into auditable analytics workflows across Sales, Finance, BI, IT, and leadership. That is directly relevant to a role that needs someone to act as the filter between business demand and technical execution while improving the credibility of core production metrics.",
        form_answer_deltas=[
            "Use the analytics leadership resume variant and the tailored Checkmk cover letter.",
            "If salary is required, stay within the existing Germany head/director default range anchored at EUR 140k-160k total compensation.",
            "Lean into sales-funnel analytics, data quality, and adoption enablement rather than pretending a pure Tableau-native background.",
        ],
        blocker_note="No blocker identified from the posting itself. This is the strongest current live role to move on next.",
        short_answer="Checkmk is attractive because the role combines BI roadmap leadership, funnel analytics, data governance, and business translation in a way that matches how I already operate across commercial analytics and executive decision support.",
    ),
    RoleConfig(
        key="pipedrive",
        company="Pipedrive",
        role="Director of Product Analytics",
        lane="Product / SaaS analytics",
        variant_key="product_saas_analytics",
        role_url="https://www.pipedrive.com/en/jobs/8933028a-9b2a-469c-9599-e449fab99029-director-of-product-analytics",
        checked_date="2026-05-12",
        location_model="Berlin hybrid",
        source="Direct company page",
        fit="Medium",
        comp_confidence="Unclear but likely above floor if true director scope; no visa sponsorship",
        work_auth_risk="Low-medium",
        stretch_type="Clear stretch on deep product analytics, experimentation, and team scale",
        action="Apply selectively after Checkmk",
        status="Prepared",
        packet_stem="Pipedrive_Director_of_Product_Analytics",
        opening_reason="it is a strong title and company-quality opportunity where product measurement, experimentation, and business growth are tightly linked. It is a stretch, but a credible one if the application leads with measurement strategy, KPI trust, and commercially useful insight rather than pretending a pure product-data-science background.",
        alignment_rows=[
            ("Measurement strategy", "Own EMEA commercial analytics across pricing, forecasting, CRM analytics, incentives, and executive business reviews."),
            ("Behavior to business link", "Translate operating signals into narratives that connect opportunity quality, conversion, revenue, margin, and customer behavior."),
            ("Self-serve analytics", "Build decision-ready workflows across Power BI, Python, SQL, machine learning, and controlled LLM systems."),
            ("Growth impact", "Delivered EUR 3M sales growth against an initial target of approximately EUR 300k and EUR 1.3M UK growth in 2023 at roughly 4x target."),
            ("Cross-functional influence", "Operate across Sales, Finance, BI, IT, and leadership to improve trust in performance data and speed of decisions."),
        ],
        strongest_alignments=[
            "Director-level title and high-quality SaaS environment.",
            "Strong overlap on measurement strategy, executive storytelling, and translating data into growth action.",
            "Clear relevance to self-serve analytics, KPI clarity, and AI-enabled efficiency language.",
        ],
        risks=[
            "The role wants 10+ years in analytics, 5+ years leading product analytics or product data science, and ideally 10+ people team scaling.",
            "Pipedrive explicitly emphasizes experimentation, retention, product-led growth, Amplitude, Optimizely, and predictive modeling in product use cases.",
            "The role is hybrid in Berlin and explicitly states no visa sponsorship or relocation assistance.",
        ],
        handling=[
            "Treat this as a selective stretch and do not oversell direct experimentation depth.",
            "Lead with product-adjacent measurement strategy, customer-behavior interpretation, and executive influence.",
            "Use the product/SaaS variant so the top third of the resume visibly matches the role language.",
        ],
        proof_points=[
            "translate complex operating data into narratives linking pipeline quality, conversion, revenue, and customer behavior",
            "built executive-ready workflows across SQL, Python, Power BI, machine learning, and controlled LLM systems",
            "delivered EUR 1.3M UK growth in 2023 at roughly 4x target",
            "delivered EUR 3M sales growth against an initial target of approximately EUR 300k",
        ],
        role_summary_tail="Strongest where a product or SaaS business needs clearer funnel visibility, more trusted KPI systems, stronger self-serve analytics, and a player-coach leader who can connect user behavior to business performance.",
        role_specific_story="Across recent roles, I have built trusted KPI systems, improved reporting integrity, and translated complex operating data into clear actions for leadership, commercial teams, and technical stakeholders. While my background is stronger in commercial and business analytics than pure in-product experimentation, the underlying leadership problem is similar: define trusted metrics, connect behavior to outcomes, and turn ambiguous questions into decision-ready insight.",
        form_answer_deltas=[
            "Use the product/SaaS analytics variant and the tailored Pipedrive cover letter.",
            "If asked why the background fits despite limited formal product-analytics titles, answer directly that the transferable value is measurement design, KPI governance, behavioral interpretation, and executive influence.",
            "Do not ignore the Berlin hybrid requirement or visa-sponsorship note when deciding whether to submit.",
        ],
        blocker_note="No formal blocker, but it is a selective stretch and should follow stronger-fit roles rather than lead the wave.",
        short_answer="Pipedrive is attractive because the role sits at the point where measurement strategy, product decisions, experimentation, and commercial growth intersect. That is close to the kind of decision-system leadership I already provide, even though my background is more commercial than pure product analytics.",
    ),
    RoleConfig(
        key="hypatos",
        company="Hypatos",
        role="Director of AI",
        lane="AI transformation / consulting",
        variant_key="ai_transformation_consulting",
        role_url="https://hypatos-gmbh.jobs.personio.com/job/2362689?language=en",
        checked_date="2026-05-12",
        location_model="Remote Germany",
        source="Direct company page",
        fit="Medium",
        comp_confidence="High enough to justify selective pursuit; role signals top-market package plus shares",
        work_auth_risk="Low",
        stretch_type="Selective stretch on deep LLM productization and enterprise automation platform depth",
        action="Apply selectively after Checkmk if AI lane remains active",
        status="Prepared",
        packet_stem="Hypatos_Director_of_AI",
        opening_reason="it is one of the better Germany-based AI leadership roles that emphasizes productization, enterprise delivery, and cross-functional execution rather than pure research. That is the exact subset of AI roles I should be selective about pursuing.",
        alignment_rows=[
            ("AI enablement", "Led European contribution to a global applied AI initiative and built training paths for safer AI adoption across real workflows."),
            ("Business translation", "Translate technical capability into practical business use cases, executive-ready communication, and governed rollout paths."),
            ("Cross-functional delivery", "Operate across Sales, Finance, BI, IT, and leadership in complex multi-country environments."),
            ("Hands-on AI depth", "Build in Python, SQL, TensorFlow, machine learning, and governed LLM workflows rather than treating AI as abstract strategy."),
            ("Commercial relevance", "Anchor AI work in forecasting, pricing, margin, CRM, and executive decision support."),
        ],
        strongest_alignments=[
            "Germany-remote model and explicit emphasis on productization rather than research alone.",
            "Clear fit to practical AI adoption, governance, stakeholder alignment, and delivery.",
            "Strong overlap on executive communication and cross-functional translation.",
        ],
        risks=[
            "The role asks for deep hands-on expertise with LLMs, agent frameworks, RAG, embeddings, orchestration frameworks, and enterprise automation domains.",
            "Hypatos is closer to an AI-native product company than a business-side AI enablement role, so the technical bar is meaningfully higher.",
        ],
        handling=[
            "Pursue only with the AI transformation variant and a cover letter that leans into productization, safe adoption, and business integration.",
            "Avoid claiming deep platform engineering or enterprise-automation implementation depth that is not well evidenced.",
            "Frame the fit around pragmatic AI delivery, workflow design, adoption, and commercial context.",
        ],
        proof_points=[
            "led European contribution to a global applied AI initiative with training and governance for safer adoption",
            "build in Python, SQL, TensorFlow, machine learning, and governed LLM workflows rather than treating AI as abstract strategy",
            "delivered EUR 1.3M UK growth in 2023 at roughly 4x target",
            "translated analytics and AI capability into executive decision support across EMEA",
        ],
        role_summary_tail="Strongest where AI needs to move from concept to governed business use cases, with clear stakeholder alignment, executive visibility, and measurable value creation rather than pure research ownership.",
        role_specific_story="Recent work has increasingly focused on practical AI adoption rather than novelty demos: identifying business-useful workflows, building governed LLM-based processes, and helping teams use AI safely and effectively within real operating environments. That is the part of the Hypatos brief that is most compelling and most aligned to my background.",
        form_answer_deltas=[
            "Use the AI transformation / consulting variant and explicitly position the profile around AI productization, workflow automation, and adoption.",
            "Be disciplined about not overstating direct RAG, orchestration-framework, or enterprise-automation implementation depth.",
        ],
        blocker_note="No formal blocker, but this should remain a selective AI application rather than displacing analytics-first roles.",
        short_answer="Hypatos is attractive because it is an AI leadership role centered on productization, enterprise delivery, and business integration rather than research in isolation.",
    ),
    RoleConfig(
        key="epam",
        company="EPAM",
        role="Director, Data Analytics Consulting",
        lane="AI transformation / consulting",
        variant_key="ai_transformation_consulting",
        role_url="https://careers.epam.com/en/vacancy/bltohj9bblkishbj98s?city=Sydney&country=Australia",
        checked_date="2026-05-12",
        location_model="Sydney, Australia",
        source="Direct company page",
        fit="Medium-high conceptually",
        comp_confidence="Likely above floor if offered, but employment model and cross-border fit need confirmation",
        work_auth_risk="Medium-high",
        stretch_type="Strong conceptual fit but Australia lane is consulting-only from Germany",
        action="Hold unless cross-border remote or contract delivery is explicitly confirmed",
        status="Prepared / Hold",
        packet_stem="EPAM_Director_Data_Analytics_Consulting",
        opening_reason="it is the clearest Australia-based role that actually matches the consulting lane: data transformation, AI implementation, client advisory, and practice growth rather than a junior or low-signal remote listing.",
        alignment_rows=[
            ("Consulting relevance", "Operate where analytics must shape executive decisions, operating models, and measurable business outcomes."),
            ("Commercial credibility", "Own EMEA commercial analytics across pricing, forecasting, margin, CRM analytics, incentives, and executive business reviews."),
            ("AI and analytics", "Build workflows across Python, SQL, Power BI, machine learning, and governed LLM systems with business value in mind."),
            ("Thought leadership", "Translate complex technical capability into executive-ready communication and practical decision systems."),
            ("Growth outcomes", "Delivered EUR 3M sales growth against an initial target of approximately EUR 300k and EUR 1.3M UK growth in 2023 at roughly 4x target."),
        ],
        strongest_alignments=[
            "Clear fit to data analytics consulting, AI implementation, client advisory, and measurable ROI language.",
            "Strong match to executive communication, solution framing, and commercial analytics leadership.",
            "Good optional lane for Australia, but only if delivery from Germany is actually workable.",
        ],
        risks=[
            "Posting is Australia-based and may assume local employment rather than cross-border consulting or contractor delivery.",
            "Role combines practice growth, presales, and enterprise consulting leadership; consulting track record must be framed carefully.",
        ],
        handling=[
            "Treat this as an opportunistic consulting lane, not a standard employment target.",
            "Use the AI transformation / consulting variant and lean into executive advisory, AI implementation, and business-outcome framing.",
            "Do not progress to submission unless location, time-zone expectations, and cross-border work model are explicitly compatible.",
        ],
        proof_points=[
            "own EMEA commercial analytics across pricing, forecasting, margin, CRM analytics, incentives, and executive business reviews",
            "build workflows across Python, SQL, Power BI, machine learning, and governed LLM systems with practical business value",
            "delivered EUR 3M sales growth against an initial target of approximately EUR 300k",
            "delivered EUR 1.3M UK growth in 2023 at roughly 4x target",
        ],
        role_summary_tail="Strongest where analytics and AI need to be translated into executive-grade transformation programs, adoption plans, and measurable business value across client environments.",
        role_specific_story="My strongest consulting-adjacent value is not novelty demos. It is helping a business define where analytics or AI should create leverage, structuring the operating logic behind that decision, and then staying close enough to the work to make sure the output is useful, auditable, and commercially relevant.",
        form_answer_deltas=[
            "Use the AI transformation / consulting variant and a cover letter that emphasizes executive advisory, AI implementation, and business outcomes.",
            "Do not apply unless the role clearly allows cross-border remote or contract delivery from Germany.",
        ],
        blocker_note="Australia roles remain contract or consulting only from Germany unless the posting explicitly confirms cross-border remote compatibility.",
        short_answer="EPAM is attractive because it is one of the few Australia-based roles in scope that actually matches a senior data and AI consulting narrative rather than generic part-time remote noise.",
    ),
]


def build_base_assets() -> None:
    for variant in VARIANTS.values():
        experience = clone_experience(variant)
        build_resume(
            BASE_VARIANTS_DIR / f"{variant.file_stem}.docx",
            variant.subtitle,
            variant.target_role,
            variant.summary,
            variant.alignment_rows,
            experience,
            EARLIER_HIGHLIGHTS,
            EDUCATION,
            SELECTED_CREDENTIALS,
            variant.stack,
        )
        build_cover_letter(
            BASE_VARIANTS_DIR / f"{variant.cover_template_stem}.docx",
            variant.cover_template_title,
            DATE_LINE,
            "[Company]\n[Role]",
            "Dear Hiring Team,",
            variant.cover_template_paragraphs,
        )

    build_recruiter_one_pager(
        BASE_VARIANTS_DIR / "Jason_Rae_Executive_Recruiter_One_Pager.docx",
        "Commercial Analytics & Applied AI Leader | Executive Recruiter One-Pager",
        VARIANTS["analytics_leadership"].summary,
        [
            "Head of Analytics",
            "Director of Analytics",
            "Head of Commercial Analytics",
            "Director of Data & Analytics",
            "Selective AI transformation leadership roles",
        ],
        [
            "EUR 3M sales growth delivered against an initial target of approximately EUR 300k",
            "EUR 1.3M UK growth in 2023 at roughly 4x target",
            "4% to 57% inside-sales turnaround across 38 markets",
            "Corrected 10+ years of FX and reporting methodology issues",
            "Current ownership of EMEA commercial analytics across pricing, forecasting, margin, CRM, incentives, and executive business reviews",
        ],
        [
            "Strong for mandates that need KPI governance, executive decision systems, and cross-functional operating change.",
            "Commercial operator to analytics architect to applied AI leader narrative is differentiated and recruiter-friendly.",
            "Works well for Germany / DACH / Europe-remote analytics leadership and selective AI transformation searches.",
        ],
        VARIANTS["analytics_leadership"].stack,
    )


def build_role_assets() -> None:
    for role in ROLES:
        variant = VARIANTS[role.variant_key]
        experience = clone_experience(variant)
        resume_name = f"Jason_Rae_{role.packet_stem}_Resume.docx"
        cover_name = f"Jason_Rae_{role.packet_stem}_Cover_Letter.docx"
        build_resume(
            APPLICATIONS_DIR / resume_name,
            variant.subtitle,
            variant.target_role,
            role_resume_summary(variant, role),
            role.alignment_rows,
            experience,
            EARLIER_HIGHLIGHTS,
            EDUCATION,
            SELECTED_CREDENTIALS,
            variant.stack,
        )
        build_cover_letter(
            APPLICATIONS_DIR / cover_name,
            f"Cover Letter | {role.role} | {role.company}",
            DATE_LINE,
            f"Hiring Team\n{role.company}",
            "Dear Hiring Team,",
            compose_role_cover_letter(variant, role),
        )
        packet = render_packet_markdown(role, variant, resume_name, cover_name)
        (RESEARCH_DIR / f"{role.packet_stem}_Application_Packet.md").write_text(
            packet,
            encoding="utf-8",
        )


def main() -> None:
    build_base_assets()
    build_role_assets()


if __name__ == "__main__":
    main()
